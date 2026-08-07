#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# Copyright © 2025-2026 Wenze Wei. All Rights Reserved.
#
# This file is part of Encre.
# The Encre project belongs to the Dunimd Team.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# You may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
# DISCLAIMER: Users must comply with applicable AI regulations.
# Non-compliance may result in service termination or legal liability.

from __future__ import annotations

"""Word document (.docx) tool.

Reads, creates and edits Microsoft Word documents, including text, headings
and tables, returning summaries or written file paths.
"""

import importlib
import json
import os
from typing import Any

from encre.tools.base import build_tool


async def _document_execute(**kwargs: Any) -> str:
    """Document execute.

    Args:
        kwargs: Description of the kwargs parameter.
    """
    action = kwargs.get("action", "")
    file_path = kwargs.get("file_path", "")
    if not file_path and action != "convert":
        return "Error: 'file_path' is required."

    if action == "read":
        return _doc_read(file_path)
    elif action == "extract_text":
        return _doc_extract_text(file_path)
    elif action == "info":
        return _doc_info(file_path)
    elif action == "create":
        content = kwargs.get("content", "")
        return _doc_create(file_path, content)
    elif action == "add_text":
        text = kwargs.get("text", "")
        return _doc_add_text(file_path, text)
    elif action == "add_table":
        data = kwargs.get("data", [])
        return _doc_add_table(file_path, data)
    elif action == "list_tables":
        return _doc_list_tables(file_path)
    elif action == "convert":
        source = kwargs.get("source", "")
        target = kwargs.get("target", "")
        return _doc_convert(source, target)
    else:
        return f"Error: Unknown action '{action}'."


def _ensure_docx() -> None:
    """Ensure docx."""
    if importlib.util.find_spec("docx") is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")


def _doc_read(file_path: str) -> str:
    """Doc read.

    Args:
        file_path: Description of the file_path parameter.
    """
    _ensure_docx()
    import docx
    if not os.path.isfile(file_path):
        return f"Error: File not found: {file_path}"
    try:
        d = docx.Document(file_path)
        parts = []
        for _i, para in enumerate(d.paragraphs):
            style = para.style.name if para.style else "Normal"
            text = para.text.strip()
            if text:
                if "Heading" in style:
                    level = style.replace("Heading ", "").split()[0] if "Heading" in style else ""
                    prefix = "#" * int(level) + " " if level.isdigit() else "## "
                    parts.append(f"{prefix}{text}")
                else:
                    parts.append(text)
        tables_info = []
        for i, table in enumerate(d.tables):
            tables_info.append(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        result = "\n".join(parts) if parts else "(empty document)"
        if tables_info:
            result += "\n\n--- Tables ---\n" + "\n".join(tables_info)
        return result
    except Exception as e:
        return f"Error reading document: {e}"


def _doc_extract_text(file_path: str) -> str:
    """Doc extract text.

    Args:
        file_path: Description of the file_path parameter.
    """
    _ensure_docx()
    import docx
    if not os.path.isfile(file_path):
        return f"Error: File not found: {file_path}"
    try:
        d = docx.Document(file_path)
        paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
        tables_text = []
        for table in d.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(cells))
        all_text = "\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)
        return all_text if all_text.strip() else "(empty document)"
    except Exception as e:
        return f"Error extracting text: {e}"


def _doc_info(file_path: str) -> str:
    """Doc info.

    Args:
        file_path: Description of the file_path parameter.
    """
    _ensure_docx()
    import docx
    if not os.path.isfile(file_path):
        return f"Error: File not found: {file_path}"
    try:
        d = docx.Document(file_path)
        core_props = d.core_properties
        para_count = len(d.paragraphs)
        table_count = len(d.tables)
        section_count = len(d.sections)
        return json.dumps({
            "file": os.path.basename(file_path),
            "size_bytes": os.path.getsize(file_path),
            "paragraphs": para_count,
            "tables": table_count,
            "sections": section_count,
            "author": str(core_props.author or ""),
            "title": str(core_props.title or ""),
            "created": str(core_props.created or ""),
            "modified": str(core_props.modified or ""),
            "last_modified_by": str(core_props.last_modified_by or ""),
        }, indent=2)
    except Exception as e:
        return f"Error reading document info: {e}"


def _doc_create(file_path: str, content: str) -> str:
    """Doc create.

    Args:
        file_path: Description of the file_path parameter.
        content: Description of the content parameter.
    """
    _ensure_docx()
    import docx

    try:
        d = docx.Document()
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("### "):
                d.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                d.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                d.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                d.add_paragraph(line[2:], style="List Bullet")
            else:
                d.add_paragraph(line)
        d.save(file_path)
        return f"Document created: {os.path.abspath(file_path)} ({len(content)} chars)"
    except Exception as e:
        return f"Error creating document: {e}"


def _doc_add_text(file_path: str, text: str) -> str:
    """Doc add text.

    Args:
        file_path: Description of the file_path parameter.
        text: Description of the text parameter.
    """
    _ensure_docx()
    import docx
    if not os.path.isfile(file_path):
        return f"Error: File not found: {file_path}"
    try:
        d = docx.Document(file_path)
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                d.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                d.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                d.add_heading(line[4:], level=3)
            else:
                d.add_paragraph(line)
        d.save(file_path)
        return f"Text added to {os.path.basename(file_path)}"
    except Exception as e:
        return f"Error adding text: {e}"


def _doc_add_table(file_path: str, data: list) -> str:
    """Doc add table.

    Args:
        file_path: Description of the file_path parameter.
        data: Description of the data parameter.
    """
    _ensure_docx()
    import docx
    if not os.path.isfile(file_path):
        return f"Error: File not found: {file_path}"
    if not data or not isinstance(data, list):
        return "Error: 'data' must be a non-empty list of rows."
    try:
        d = docx.Document(file_path)
        headers = data[0] if isinstance(data[0], list) else []
        rows_data = data[1:] if headers else data
        table = d.add_table(rows=1 + len(rows_data), cols=len(headers) if headers else len(rows_data[0]) if rows_data else 1)
        table.style = "Table Grid"
        if headers:
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = str(h)
            for r_idx, row in enumerate(rows_data):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx + 1].cells[c_idx].text = str(val)
        else:
            for r_idx, row in enumerate(rows_data):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = str(val)
        d.save(file_path)
        return f"Table added to {os.path.basename(file_path)}"
    except Exception as e:
        return f"Error adding table: {e}"


def _doc_list_tables(file_path: str) -> str:
    """Doc list tables.

    Args:
        file_path: Description of the file_path parameter.
    """
    _ensure_docx()
    import docx
    if not os.path.isfile(file_path):
        return f"Error: File not found: {file_path}"
    try:
        d = docx.Document(file_path)
        if not d.tables:
            return "No tables in document."
        out = []
        for i, table in enumerate(d.tables):
            rows_data = []
            for row in table.rows[:5]:
                rows_data.append([cell.text.strip()[:30] for cell in row.cells])
            out.append(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
            for r in rows_data:
                out.append("  " + " | ".join(r))
            if len(table.rows) > 5:
                out.append(f"  ... ({len(table.rows) - 5} more rows)")
        return "\n".join(out)
    except Exception as e:
        return f"Error listing tables: {e}"


def _doc_convert(source: str, target: str) -> str:
    """Doc convert.

    Args:
        source: Description of the source parameter.
        target: Description of the target parameter.
    """
    if not source or not target:
        return "Error: 'source' and 'target' paths required for convert."
    if not os.path.isfile(source):
        return f"Error: Source file not found: {source}"

    _ensure_docx()
    import docx

    ext = os.path.splitext(target)[1].lower()
    try:
        d = docx.Document(source)
        text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        if ext == ".txt":
            with open(target, "w", encoding="utf-8") as f:
                f.write(text)
        elif ext == ".md":
            with open(target, "w", encoding="utf-8") as f:
                for p in d.paragraphs:
                    t = p.text.strip()
                    if not t:
                        continue
                    if "Heading" in (p.style.name or ""):
                        level = p.style.name.replace("Heading ", "")
                        f.write("#" * int(level) + " " + t + "\n\n")
                    else:
                        f.write(t + "\n\n")
        else:
            return f"Error: Unsupported target format '{ext}'. Supported: .txt, .md"
        return f"Converted: {source} -> {target}"
    except Exception as e:
        return f"Error converting document: {e}"


EncreDocumentTool = build_tool(
    name="document",
    description=(
        "Read, create, edit, and convert Microsoft Word (.docx) documents via "
        "python-docx, including paragraphs, headings, and tables. "
        "Use this for actions such as `read` (formatted dump), `extract_text` (plain "
        "text), `info` (metadata), `create`/`add_text`/`add_table` (write), "
        "`list_tables`, or `convert` to .txt/.md. "
        "Do NOT use this for PDFs (use the pdf tool), spreadsheets (use spreadsheet), "
        "or presentations (use presentation); and avoid it for raw .odt editing. "
        "Tips: in `create`/`add_text`, use markdown-style prefixes (#, ##, ###, - or *) "
        "to trigger headings and bullet lists; pass `data` as a 2D array for tables. "
        "Pitfalls: requires python-docx installed; `convert` only supports .txt and .md."
    ),
    input_schema={
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": ["read", "extract_text", "info", "create", "add_text", "add_table", "list_tables", "convert"],
                "description": "Document operation: read (formatted), extract_text (plain), info (metadata), create (new file), add_text (append), add_table (append table), list_tables (preview tables), convert (docx->txt/md).",
            },
            "file_path": {
                "type": "string",
                "description": "Path to the .docx file to read or modify; required for all actions except convert.",
            },
            "content": {
                "type": "string",
                "description": "Body for the create action; supports markdown-style headings (#, ##, ###) and bullet lists (- or *).",
            },
            "text": {
                "type": "string",
                "description": "Text to append via add_text; lines starting with #, ##, ### become headings.",
            },
            "data": {
                "type": "array",
                "items": {"type": "array"},
                "description": "Table rows for add_table; first row is treated as headers when present (e.g. [[\"Name\",\"Age\"],[\"Ada\",36]]).",
            },
            "source": {
                "type": "string",
                "description": "Source .docx path for the convert action.",
            },
            "target": {
                "type": "string",
                "description": "Destination path for convert; extension must be .txt or .md.",
            },
        },
        "required": ["action"],
    },
    execute=_document_execute,
    intents=["data", "general", "research"],
    category="docs",
    semantic_type="read",
    is_concurrency_safe=lambda data: data.get("action") in ("read", "extract_text", "info", "list_tables"),
    is_destructive=lambda args: args.get("action", "") in ("create", "add_text", "add_table", "convert"),
)
